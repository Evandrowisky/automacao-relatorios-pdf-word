"""Geração de relatórios Word a partir de um modelo."""

import logging
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import Inches


LOGGER = logging.getLogger(__name__)


def generate_word_report(
    template_path: Path,
    output_path: Path,
    data: dict[str, str],
    image_paths: list[Path] | None = None,
) -> Path:
    """Gera um relatório Word substituindo placeholders do modelo."""

    if not template_path.exists():
        raise FileNotFoundError(f"Modelo Word não encontrado: {template_path}")

    LOGGER.info("Gerando relatório Word: %s", output_path)
    document = Document(template_path)
    replace_placeholders(document, data)

    if image_paths:
        append_images(document, image_paths)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    LOGGER.info("Relatório Word gerado com sucesso.")

    return output_path


def replace_placeholders(document: DocumentObject, data: dict[str, str]) -> None:
    """Substitui placeholders no formato {{campo}} pelo valor correspondente."""

    replacements = {f"{{{{{key}}}}}": value for key, value in data.items()}

    for paragraph in document.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)


def replace_text_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """Substitui textos em um parágrafo preservando a estrutura principal."""

    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)


def append_images(document: DocumentObject, image_paths: list[Path]) -> None:
    """Adiciona imagens ao final do relatório."""

    document.add_page_break()
    document.add_heading("Imagens", level=2)

    for image_path in image_paths:
        document.add_paragraph(image_path.name)
        document.add_picture(str(image_path), width=Inches(5.5))
